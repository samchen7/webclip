import os
import re
import json
import logging
from pathlib import Path
from typing import List
from docx import Document

log = logging.getLogger(__name__)

def safe_name(s: str) -> str:
    s = re.sub(r"[^a-zA-Z0-9._-]+", "_", s)
    return s[:100]

class FileManager:
    def __init__(self, base_dir: str):
        self.base = Path(base_dir)
        self.base.mkdir(parents=True, exist_ok=True)

    def save_text(self, url: str, title: str, text: str) -> str:
        name = safe_name(url) + ".md"
        p = self.base / name
        with open(p, "w", encoding="utf-8") as f:
            f.write(f"# {title}\n\n")
            f.write(text)
        return str(p)

    def save_images(self, url: str, image_paths: List[str]) -> List[str]:
        # 假设已经是本地路径；若是下载的远程图片，可在此实现下载
        return image_paths

    def save_partial_docx(self, url: str, title: str, h1: list[str], h2: list[str], paragraphs: list[str], note: str="") -> str:
        name = safe_name(url) + ".docx"
        p = self.base / name
        doc = Document()
        doc.add_heading(title or url, 0)
        if note:
            doc.add_paragraph(f"[NOTE] {note}")
        for t in h1:
            doc.add_heading(t, level=1)
        for t in h2:
            doc.add_heading(t, level=2)
        for para in paragraphs:
            doc.add_paragraph(para)
        doc.save(p)
        return str(p)

    def save_json(self, url: str, payload: dict) -> str:
        name = safe_name(url) + ".json"
        p = self.base / name
        with open(p, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        return str(p) 